# output_formatter.py -- Phase 8
# Converts a DocumentExtraction dict into PDF / DOCX / XLSX / CSV / JSON files.
import json
from pathlib import Path

import pandas as pd
from fpdf import FPDF
from docx import Document
from openpyxl import Workbook
import matplotlib


# ── Unicode font setup for fpdf2 ──────────────────────────────────────────
# fpdf2 supports TrueType fonts with full Unicode. We use matplotlib's bundled
# DejaVuSans so we don't need to ship font files ourselves.
_FONT_DIR = Path(matplotlib.get_data_path()) / "fonts/ttf"

DEJAVU_REGULAR = str(_FONT_DIR / "DejaVuSans.ttf")
DEJAVU_BOLD    = str(_FONT_DIR / "DejaVuSans-Bold.ttf")
DEJAVU_ITALIC  = str(_FONT_DIR / "DejaVuSans-Oblique.ttf")  # oblique = italic in DejaVu naming

# Verify fonts exist (graceful fallback message if not)
for _f in (DEJAVU_REGULAR, DEJAVU_BOLD, DEJAVU_ITALIC):
    if not Path(_f).exists():
        raise FileNotFoundError(
            f"Required font not found: {_f}\n"
            "Install matplotlib or place DejaVu TTF fonts in a 'fonts/' folder."
        )


def _register_fonts(pdf: FPDF):
    """Add Unicode fonts to the PDF object once per document."""
    pdf.add_font("DejaVu", "",      DEJAVU_REGULAR, uni=True)
    pdf.add_font("DejaVu", "B",     DEJAVU_BOLD,    uni=True)
    pdf.add_font("DejaVu", "I",     DEJAVU_ITALIC,  uni=True)
    # Alias bold-italic if needed later
    pdf.add_font("DejaVu", "BI",    DEJAVU_BOLD,    uni=True)  # fallback to bold


def _line_items_df(data):
    items = data.get("line_items", [])
    if not items:
        items = [{"description": "", "quantity": 0, "unit_price": 0.0, "total": 0.0}]
    return pd.DataFrame(items)


def save_json(data, out_path):
    Path(out_path).write_text(json.dumps(data, indent=2), encoding="utf-8")
    return str(out_path)


def save_csv(data, out_path):
    # Flat one-row summary + line items appended below
    flat = {
        "vendor_name":   data["vendor"]["name"],
        "document_type": data["document"]["type"],
        "date":          data["document"]["date"],
        "subtotal":      data["totals"]["subtotal"],
        "tax":           data["totals"]["tax"],
        "discount":      data["totals"]["discount"],
        "total":         data["totals"]["total"],
        "amount_paid":   data["payment"]["amount_paid"],
        "confidence":    data["extraction_meta"]["overall_confidence"],
    }
    df = pd.DataFrame([flat])
    df.to_csv(out_path, index=False)
    return str(out_path)


def save_xlsx(data, out_path):
    wb = Workbook()

    # Sheet 1: summary
    ws = wb.active
    ws.title = "Summary"
    rows = [
        ("Vendor",      data["vendor"]["name"]),
        ("Type",        data["document"]["type"]),
        ("Date",        data["document"]["date"]),
        ("Subtotal",    data["totals"]["subtotal"]),
        ("Tax",         data["totals"]["tax"]),
        ("Discount",    data["totals"]["discount"]),
        ("Total",       data["totals"]["total"]),
        ("Amount Paid", data["payment"]["amount_paid"]),
        ("Confidence",  data["extraction_meta"]["overall_confidence"]),
        ("Time (ms)",   data["extraction_meta"]["processing_time_ms"]),
    ]
    for r in rows:
        ws.append(r)

    # Sheet 2: line items
    ws2 = wb.create_sheet("Line Items")
    ws2.append(["Description", "Quantity", "Unit Price", "Total"])
    for item in data.get("line_items", []):
        ws2.append([item["description"], item["quantity"],
                    item["unit_price"], item["total"]])

    wb.save(out_path)
    return str(out_path)


def save_docx(data, out_path):
    doc = Document()
    doc.add_heading("Document Extraction Report", level=1)

    doc.add_heading("Vendor", level=2)
    doc.add_paragraph(f'Name: {data["vendor"]["name"]}')
    doc.add_paragraph(f'Address: {data["vendor"]["address"]}')

    doc.add_heading("Document", level=2)
    doc.add_paragraph(f'Type: {data["document"]["type"]}  |  Date: {data["document"]["date"]}')

    doc.add_heading("Line Items", level=2)
    items = data.get("line_items", [])
    table = doc.add_table(rows=1, cols=4)
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    hdr[0].text, hdr[1].text, hdr[2].text, hdr[3].text = "Description", "Qty", "Unit Price", "Total"
    for item in items:
        row = table.add_row().cells
        row[0].text = str(item["description"])
        row[1].text = str(item["quantity"])
        row[2].text = str(item["unit_price"])
        row[3].text = str(item["total"])

    doc.add_heading("Totals", level=2)
    t = data["totals"]
    doc.add_paragraph(f'Subtotal: {t["subtotal"]}  |  Tax: {t["tax"]}  |  '
                      f'Discount: {t["discount"]}  |  TOTAL: {t["total"]}')

    meta = data["extraction_meta"]
    doc.add_paragraph(f'Confidence: {meta["overall_confidence"]}  |  '
                      f'Time: {meta["processing_time_ms"]}ms  |  '
                      f'Low-confidence fields: {", ".join(meta["low_confidence_fields"]) or "none"}')

    doc.save(out_path)
    return str(out_path)


def save_pdf(data, out_path, overlay_path=None):
    pdf = FPDF()
    _register_fonts(pdf)   # <-- register Unicode fonts

    pdf.add_page()
    pdf.set_font("DejaVu", "B", 16)
    pdf.cell(0, 10, "Document Extraction Report", ln=True)

    pdf.set_font("DejaVu", "", 11)
    pdf.ln(4)
    pdf.cell(0, 8, f'Vendor: {data["vendor"]["name"]}', ln=True)
    pdf.cell(0, 8, f'Type: {data["document"]["type"]}   Date: {data["document"]["date"]}', ln=True)
    pdf.ln(2)

    # Line items table
    pdf.set_font("DejaVu", "B", 11)
    pdf.cell(90, 8, "Description", border=1)
    pdf.cell(25, 8, "Qty",        border=1)
    pdf.cell(35, 8, "Unit Price", border=1)
    pdf.cell(35, 8, "Total",      border=1, ln=True)
    pdf.set_font("DejaVu", "", 10)
    for item in data.get("line_items", []):
        # Unicode-safe: no more latin-1 limitation
        pdf.cell(90, 8, str(item["description"])[:45], border=1)
        pdf.cell(25, 8, str(item["quantity"]),   border=1)
        pdf.cell(35, 8, str(item["unit_price"]), border=1)
        pdf.cell(35, 8, str(item["total"]),      border=1, ln=True)

    pdf.ln(4)
    t = data["totals"]
    pdf.set_font("DejaVu", "B", 12)
    pdf.cell(0, 8, f'Subtotal: {t["subtotal"]}   Tax: {t["tax"]}   TOTAL: {t["total"]}', ln=True)

    meta = data["extraction_meta"]
    pdf.set_font("DejaVu", "I", 9)
    pdf.cell(0, 8, f'Confidence: {meta["overall_confidence"]}  |  '
                   f'Pipeline: {meta["model_path"]}  |  Time: {meta["processing_time_ms"]}ms', ln=True)

    # Page 2: annotated overlay image
    if overlay_path and Path(overlay_path).exists():
        pdf.add_page()
        pdf.set_font("DejaVu", "B", 14)
        pdf.cell(0, 10, "Annotated Overlay", ln=True)
        pdf.image(str(overlay_path), x=10, y=25, w=180)

    pdf.output(str(out_path))
    return str(out_path)


def export_all(data, out_dir, base_name, overlay_path=None):
    # Generate all 5 formats, return list of file paths
    out_dir = Path(out_dir)
    out_dir.mkdir(parents=True, exist_ok=True)
    files = []
    files.append(save_json(data, out_dir / f"{base_name}.json"))
    files.append(save_csv(data,  out_dir / f"{base_name}.csv"))
    files.append(save_xlsx(data, out_dir / f"{base_name}.xlsx"))
    files.append(save_docx(data, out_dir / f"{base_name}.docx"))
    files.append(save_pdf(data,  out_dir / f"{base_name}.pdf", overlay_path))
    return files